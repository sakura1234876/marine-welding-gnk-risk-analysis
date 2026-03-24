# 1.Convert the Word file to a TXT file
from docx import Document

# WORD DOCUMENT EXTRACTION

def read_word_file(word_path):
    """
    Reads a Microsoft Word (.docx) document and extracts all text content.
    
    Args:
        word_path (str): The path to the source Word document.
        
    Returns:
        str: A single string containing all non-empty paragraphs.
    """
    # Initialize the Word document reader
    doc = Document(word_path)
    full_content = []
    
    # Iterate through paragraphs and skip empty lines
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_content.append(paragraph.text.strip())
            
    # Join all lines into a single block of text for processing
    return "\n".join(full_content)
# DATA SEGMENTATION & TXT EXPORT
def convert_word_content_to_txt(word_content, output_txt_path):
    """
    Segments raw text into individual cases and saves them to a TXT file.
    
    This function splits the text using the keyword 'Case' (案例) and 
    separates each valid case with a unique '##########' delimiter.

    Args:
        word_content (str): Raw string extracted from the document.
        output_txt_path (str): Path where the formatted TXT will be saved.
    """
    # Split content by the keyword 'Case'
    case_segments = word_content.split("Case")
    valid_cases = []
    
    for segment in case_segments:
        segment = segment.strip()
        if not segment:
            continue
            
        # Ensure the segment corresponds to a numbered case (e.g., 'Case 1')
        if segment[0].isdigit():
            valid_case = f"Case{segment}"
            valid_cases.append(valid_case)
    
    # Create a visual separator for better readability and chunking
    separator = "\n" + "#" * 10 + "\n\n"
    txt_content = separator.join(valid_cases)
    
    # Write the formatted output to a local text file
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
    
    # Log progress for monitoring
    print(f"Conversion Complete! File saved to: {output_txt_path}")
    print(f"Processed {len(valid_cases)} valid cases.")

# EXECUTION FLOW 
# Placeholder for the source Word document (e.g., "Quality_Cases.docx")
SOURCE_FILE = "<YOUR_INPUT_WORD_FILE_PATH>"
# Placeholder for the destination Text file (e.g., "Formatted_Cases.txt")
OUTPUT_FILE = "<YOUR_OUTPUT_TXT_FILE_PATH>"
# Step 1: Extract content from the Word file
raw_text = read_word_file(SOURCE_FILE)
# Step 2: Pass extracted content to the converter to generate the TXT
convert_word_content_to_txt(raw_text, OUTPUT_FILE)

#--------------------------------
# 2. Convert the TXT file to a csv file
#--------------------------------
import os
import csv
import re
from docx import Document
# TEXT EXTRACTION (WORD TO STRING)
def read_word_file(word_path):
    """
    Reads a Word document and returns its content as a cleaned string.
    
    Args:
        word_path (str): Path to the source .docx file.
    Returns:
        str: Consolidated text with empty lines filtered out.
    """
    try:
        doc = Document(word_path)
        full_content = []
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:  # Filter out empty lines
                full_content.append(para_text)
        return "\n".join(full_content)
    except Exception as e:
        print(f"Error reading Word file: {e}")
        return ""
# DATA PARSING (REGEX EXTRACTION)
def extract_case_data(word_content):
    """
    Extracts specific fields: Occurrence Time, Stage, Phenomenon, and Root Cause Analysis.
    
    Note: The split keyword 'Case' and field headers are now localized to English.
    """
    # Split the content by the keyword 'Case' to isolate individual records
    case_segments = [seg.strip() for seg in word_content.split("Case") if seg.strip()]
    all_case_data = []
    
    # Regex patterns using Lookahead (?=...) to isolate text between headers
    # Supports both full-width (：) and half-width (:) colons
    pattern_time = re.compile(r"1\. Occurrence Time[\s:：]+(.*?)(?=2\. Occurrence Stage)", re.DOTALL)
    pattern_stage = re.compile(r"2\. Occurrence Stage[\s:：]+(.*?)(?=3\. Problem Phenomenon)", re.DOTALL)
    pattern_phenomenon = re.compile(r"3\. Problem Phenomenon[\s:：]+(.*?)(?=4\. Root Cause Analysis)", re.DOTALL)
    pattern_reason = re.compile(r"4\. Root Cause Analysis[\s:：]+(.*)", re.DOTALL)
    
    for seg in case_segments:
        # Search for matches within the current segment
        time_match = re.search(pattern_time, seg)
        stage_match = re.search(pattern_stage, seg)
        phen_match = re.search(pattern_phenomenon, seg)
        reason_match = re.search(pattern_reason, seg)
        
        def clean_text(match_obj):
            """Helper to strip whitespace and normalize internal spacing."""
            if match_obj:
                text = match_obj.group(1).strip()
                # Replace multiple spaces/newlines with a single space for CSV compatibility
                return re.sub(r"\s+", " ", text)
            return ""
        
        # Consolidate extracted data into a row
        case_info = [
            clean_text(time_match),
            clean_text(stage_match),
            clean_text(phen_match),
            clean_text(reason_match)
        ]
        all_case_data.append(case_info)
    
    return all_case_data
# 3. STORAGE (CSV EXPORT)
def save_to_csv(case_data, csv_path):
    """
    Writes the structured data into a CSV file with localized headers.
    """
    headers = [
        "Occurrence Time", 
        "Occurrence Stage", 
        "Problem Phenomenon", 
        "Root Cause Analysis"
    ]
    
    # Using 'utf-8-sig' ensures compatibility with Excel
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(case_data)
        
    print(f"CSV generated successfully! Extracted {len(case_data)} records.")
    print(f"Output Path: {csv_path}")

if __name__ == "__main__":
    # --- CONFIGURATION PLACEHOLDERS ---
    # Replace these with your actual file paths
    INPUT_WORD_FILE = "<YOUR_INPUT_WORD_FILE_PATH>"
    OUTPUT_CSV_FILE = "<YOUR_OUTPUT_CSV_FILE_PATH>"
    # Process Flow: Read -> Extract -> Save
    if os.path.exists(INPUT_WORD_FILE):
        content = read_word_file(INPUT_WORD_FILE)
        data = extract_case_data(content)
        save_to_csv(data, OUTPUT_CSV_FILE)
    else:
        print(f"Error: Input file '{INPUT_WORD_FILE}' not found.")

#--------------------------------
# 3. Factor classification
#--------------------------------
# Main factor classification
import requests
import json
import pandas as pd
from tqdm import tqdm
from retry import retry

# Configuration & Constants ---
OLLAMA_BASE = "http://<ollama host>:11434"
CHAT_MODEL = "qwen3:30b-a3b-Instruct" #model name

# The System Prompt defines the AI's persona as a marine welding expert 
# and provides the 5M1E (Human, Equipment, Material, Technology, Environment, Management) framework.
SYSTEM_PROMPT = (
    "You are a professional marine welding risk analysis expert. Your task is to analyze "
    "the provided 'Root Cause Analysis' text and classify identified risk factors into "
    "the following six predefined categories:\n\n"
    "1. **Human**: Factors related to operators (e.g., skill, experience, responsibility, habits, fatigue).\n"
    "   * Keywords: Lack of responsibility, insufficient skill, uncertified, illegal operation, new employee, poor training.\n\n"
    "2. **Equipment**: Factors related to tools and equipment (e.g., performance, maintenance, functionality).\n"
    "   * Keywords: Damaged equipment, abnormal function, poor maintenance, broken cables, faulty indicators, missing parts.\n\n"
    "3. **Material**: Factors related to raw materials and consumables (e.g., specifications, quality, storage).\n"
    "   * Keywords: Wrong material, unqualified raw materials, rusty wire, damp electrodes, impurities, wrong wire type.\n\n"
    "4. **Technology**: Factors related to work methods, standards, and design (e.g., procedures, blueprints, technical requirements).\n"
    "   * Keywords: Violated process requirements, blueprint deviation, design defect, improper groove, wrong welding sequence, no preheating.\n\n"
    "5. **Environment**: Physical workplace factors (e.g., weather, temperature, humidity, lighting, space).\n"
    "   * Keywords: Rain/snow, high humidity, low temperature, strong wind, narrow space, high construction difficulty.\n\n"
    "6. **Management**: Factors related to organization and supervision (e.g., planning, oversight, communication, workflow control).\n"
    "   * Keywords: No inspection, lack of supervision, poor duty performance, failed briefing, insufficient manpower, skipped procedures.\n\n"
    "**Your Task:**\n"
    "- Read the 'Root Cause Analysis' text carefully.\n"
    "- Extract all explicitly mentioned risk factors.\n"
    "- Categorize each factor into one of the six categories above.\n"
    "- Output ONLY a strict JSON object without any extra explanation.\n\n"
    "**Output Format:**\n"
    "```json\n"
    "{\n"
    "  \"Human\": [\"factor 1\", \"factor 2\", ...],\n"
    "  \"Equipment\": [...],\n"
    "  \"Material\": [...],\n"
    "  \"Technology\": [...],\n"
    "  \"Environment\": [...],\n"
    "  \"Management\": [...]\n"
    "}\n"
    "```"
)
#  Core Classification Logic ---
@retry(requests.exceptions.RequestException, tries=3, delay=1)
def classify_risk_factors(reason_text):
    """
    Sends text to the LLM to extract and classify risk factors into JSON.
    Includes error handling and content cleaning for robust JSON parsing.
    """
    payload = {
        "model": CHAT_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": reason_text}
        ],
        "temperature": 0.1,  # Low temperature for consistent, factual output
        "stream": False
    }

    try:
        resp = requests.post(f"{OLLAMA_BASE}/api/chat", json=payload)
        resp.raise_for_status()
        
        content = resp.json()["message"]["content"]
        
        # Strip Markdown code blocks if present
        if content.startswith("```json") and content.endswith("```"):
            content = content[7:-3].strip()

        try:
            return json.loads(content)
        except json.JSONDecodeError as json_err:
            print(f"JSON Parsing Error: {json_err} | Raw Content: {content[:100]}...")
            return {cat: [] for cat in ["Human", "Equipment", "Material", "Technology", "Environment", "Management"]}

    except Exception as e:
        print(f"Error processing text: {str(e)}")
        return {cat: [] for cat in ["Human", "Equipment", "Material", "Technology", "Environment", "Management"]}

#  Batch Processing Execution ---

def run_analysis(input_csv, output_xlsx):
    """
    Reads the CSV file, processes each row via LLM, and saves the categorized results.
    """
    # Load dataset (using GBK encoding as common for legacy Chinese Excel exports)
    df = pd.read_csv(input_csv, encoding='GBK')
    
    # Initialize new result columns
    df["Risk_Factors"] = ""
    df["Factor_Categories"] = ""

    for index, row in tqdm(df.iterrows(), total=len(df), desc="Processing Cases"):
        reason_text = str(row["Problem_Reason_Analysis"]) # Ensure header matches your CSV
        
        if not reason_text or reason_text.strip().lower() == "nan":
            continue

        # Get classified data from AI
        risk_data = classify_risk_factors(reason_text)
        
        all_factors = []
        all_categories = []

        # Aggregate factors and their corresponding category labels
        for category, factors in risk_data.items():
            if factors:
                all_factors.extend(factors)
                all_categories.extend([category] * len(factors))

        # Deduplicate results for the current row
        unique_factors = list(dict.fromkeys(all_factors))
        unique_categories = list(dict.fromkeys(all_categories))

        # Join results into strings for Excel storage
        df.at[index, "Risk_Factors"] = "; ".join(unique_factors)
        df.at[index, "Factor_Categories"] = "; ".join(unique_categories)

    # Save to final Excel report
    df.to_excel(output_xlsx, index=False)
    print(f"\nAnalysis complete. Results saved to: {output_xlsx}")

if __name__ == "__main__":
    # --- Placeholders for Input/Output ---
    INPUT_FILE = "<YOUR_INPUT_CASE_DATA>.csv"
    OUTPUT_FILE = "<YOUR_FINAL_RISK_REPORT>.xlsx"
    run_analysis(INPUT_FILE, OUTPUT_FILE)

#--------------------------------
# 4. Sub_Factor classification
#--------------------------------
import pandas as pd
import time
import requests
import json
#  Configuration & Constants ---
OLLAMA_API_URL = "http://<ollama host>/api/generate" 
MODEL_NAME = "qwen3:30b-a3b-Instruct" 
# Detailed prompt instructing the AI to categorize welding issues into a 2-tier factor system
CLASSIFICATION_RULES_PROMPT = """
You are a professional marine welding quality inspector.
Categorize the provided "Factor" text into specific **Secondary Sub-factors** within the 6 primary categories. 
Each input must be assigned to **exactly one** most relevant Secondary Sub-factor. Do NOT use the primary category names (e.g., "Human Factors").

Primary Categories and their Secondary Sub-factors:

1. Human Factors:
   - Skill Level: Insufficient skill, poor training, new employees, unable to read blueprints, uncertified.
   - Responsibility & Attitude: Lack of responsibility, weak quality awareness, careless work, blind construction.
   - Operating Procedures: Illegal operations, violating blueprints/processes, improper handling, wrong welding sequence.

2. Equipment Factors:
   - Equipment Status: Abnormal equipment, damaged functions, missing parts, poor maintenance, broken cables.
   - Tool Usage: Improper tool use, unqualified tools, insufficient protection, welder failure.
   - Performance Indicators: Abnormal voltage/current, excessive speed, insufficient heating temperature.

3. Material Factors:
   - Material Quality: Unqualified materials, raw materials not meeting requirements, excess hardness, unauthorized substitution.
   - Consumable Management: Rusty wire, damp electrodes, wrong wire type used.
   - Storage Conditions: Improper temporary storage, insufficient oven temperature, damp/rusty consumables.

4. Technology Factors:
   - Process Design: Design defects, incomplete processes, groove issues, lack of targeted process plans.
   - Process Execution: Improper parameters, wrong methods, chaotic sequence, insufficient pretreatment, fast welding.
   - Technical Standards: Unclear standards, mismatched specs, insufficient weld leg, excessive gap.

5. Environment Factors:
   - Weather Conditions: Impact of rain, snow, low temperature, high humidity, monsoon season.
   - Construction Environment: Narrow space, high difficulty, poor conditions, strong wind, poor lighting.

6. Management Factors:
   - Supervision & Inspection: Inadequate inspection, missing mutual checks, lack of oversight, failure to detect issues.
   - Systems & Procedures: Imperfect systems, failure to implement requirements, non-standard workflows, poor dispatching.
   - Training & Briefing: Insufficient training, unclear briefing, requirements not communicated, lack of pre-job training.

[Classification Rules]
- Return ONLY these Secondary Sub-factor names: Skill Level, Responsibility & Attitude, Operating Procedures, Equipment Status, Tool Usage, Performance Indicators, Material Quality, Consumable Management, Storage Conditions, Process Design, Process Execution, Technical Standards, Weather Conditions, Construction Environment, Supervision & Inspection, Systems & Procedures, Training & Briefing.
- DO NOT return primary category names.
- If no match is found, classify as "Unclassified".
- Return result as a strict JSON object. No explanations or extra text.

Return format:
{
  "secondary_factors": ["Sub-factor1", "Sub-factor2", ...]
}
"""

# Classification Function ---
def classify_factor_with_ollama(factor_string):
    """
    Classifies a string of factor descriptions using the local Ollama LLM.
    
    Args:
        factor_string (str): The raw text describing a welding issue or path.
    Returns:
        str: Semi-colon separated secondary sub-factors.
    """
    prompt = f"{CLASSIFICATION_RULES_PROMPT}\n\nFactor: {factor_string}"
    data = {
        "model": MODEL_NAME, 
        "prompt": prompt,
        "stream": False, 
        "options": {
            "temperature": 0, # Low temperature for deterministic output
            "num_predict": 500 
        }
    }

    try:
        response = requests.post(OLLAMA_API_URL, json=data)
        response.raise_for_status() 
        response_json = response.json()
        response_content = response_json.get("response", "").strip()
        
        # Parse JSON from LLM response
        result = json.loads(response_content)
        subcategories = result.get("secondary_factors", [])
        
        # Filter out 'Unclassified' and join results
        matched = [cat for cat in subcategories if cat != "Unclassified"]
        return '; '.join(matched) if matched else ''
        
    except requests.exceptions.RequestException as e:
        print(f"Network error calling Ollama API for: {factor_string}. Error: {e}")
        return ''
    except json.JSONDecodeError as e:
        print(f"Failed to parse JSON for: {factor_string}. Response: {response_content[:100]}... Error: {e}")
        return '' 
    except Exception as e:
        print(f"Unknown error during classification for: {factor_string}. Error: {e}")
        return '' 

# Main Execution Flow ---

# Configuration for input/output files (Replace with your actual paths)
INPUT_FILE = "<main_factor>.xlsx"
OUTPUT_FILE = "Classified.xlsx"

# Load the source Excel sheet
try:
    df = pd.read_excel(INPUT_FILE, sheet_name='Sheet1')
    results = []

    # Iterate through rows to classify the 'max_path' column
    for index, row in df.iterrows():
        factor = row['max_path']
        print(f"Processing row {index + 1}: {factor}")
        
        classified_result = classify_factor_with_ollama(factor)
        results.append(classified_result)
        
        # Sleep to manage local LLM load/rate limits
        time.sleep(2) 

    # Save results to a new column and export
    df['Secondary_Factors'] = results
    df.to_excel(output_FILE, index=False, sheet_name='Sheet1')

    print(f"\nClassification complete. Results saved to: {OUTPUT_FILE}")

except Exception as e:
    print(f"Critical error during execution: {e}")
